import os
import tempfile
from docx import Document
from docx2pdf import convert
from pdf2docx import Converter

class StrikethroughRedactor:
    """
    Class to redact strikethrough text from DOCX or PDF files
    and save the cleaned documents as PDFs.
    """

    def __init__(self, input_dir: str, output_dir: str):
        """
        Initialize directories.
        """
        self.input_dir = input_dir
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def redact_strikethrough(self, doc: Document) -> None:
        """
        Remove strikethrough text from a Document object.
        """
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.strike:
                    run.text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.strike:
                                run.text = ''

    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str) -> None:
        """
        Convert a PDF file to a DOCX file.
        """
        converter = Converter(pdf_path)
        converter.convert(docx_path)
        converter.close()

    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """
        Convert a DOCX file to a PDF file.
        """
        convert(docx_path, pdf_path)

    def has_strikethrough(self, doc: Document) -> bool:
        """
        Check if a Document has any strikethrough text.
        """
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.strike:
                    return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.strike:
                                return True
        return False

    def process_docx(self, input_docx_path: str, output_pdf_path: str) -> None:
        """
        Process a DOCX file: remove strikethrough and save as PDF.
        """
        doc = Document(input_docx_path)
        self.redact_strikethrough(doc)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            temp_docx_path = tmp_file.name

        try:
            doc.save(temp_docx_path)
            self.convert_docx_to_pdf(temp_docx_path, output_pdf_path)
        finally:
            os.remove(temp_docx_path)

    def process_pdf(self, input_pdf_path: str, output_pdf_path: str) -> None:
        """
        Process a PDF file: convert to DOCX, remove strikethrough, and save back as PDF.
        """
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx_file:
            temp_docx_path = tmp_docx_file.name

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_clean_docx_file:
            cleaned_docx_path = tmp_clean_docx_file.name

        try:
            self.convert_pdf_to_docx(input_pdf_path, temp_docx_path)
            doc = Document(temp_docx_path)

            if self.has_strikethrough(doc):
                print(f"Strikethrough detected. Cleaning: {input_pdf_path}")
                self.redact_strikethrough(doc)
                doc.save(cleaned_docx_path)
                self.convert_docx_to_pdf(cleaned_docx_path, output_pdf_path)
            else:
                print(f"No strikethrough found. Copying original: {input_pdf_path}")
                self.convert_docx_to_pdf(temp_docx_path, output_pdf_path)

        finally:
            for temp_file in [temp_docx_path, cleaned_docx_path]:
                if os.path.exists(temp_file):
                    os.remove(temp_file)

    def process_directory(self) -> None:
        """
        Process all DOCX and PDF files in the input directory.
        """
        for filename in os.listdir(self.input_dir):
            input_path = os.path.join(self.input_dir, filename)
            output_filename = os.path.splitext(filename)[0] + ".pdf"
            output_path = os.path.join(self.output_dir, output_filename)

            if filename.lower().endswith('.docx'):
                print(f"Processing DOCX: {filename}")
                self.process_docx(input_path, output_path)
                print(f"Saved cleaned PDF: {output_path}")

            elif filename.lower().endswith('.pdf'):
                print(f"Processing PDF: {filename}")
                self.process_pdf(input_path, output_path)
                print(f"Saved cleaned PDF: {output_path}")

if __name__ == "__main__":
    input_directory_path = r"your\input\directory\path"
    output_directory_path = r"your\output\directory\path"

    redactor = StrikethroughRedactor(input_directory_path, output_directory_path)
    redactor.process_directory()
